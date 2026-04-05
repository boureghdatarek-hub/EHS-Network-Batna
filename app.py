import streamlit as st
import pandas as pd
from datetime import datetime
import os
import base64
from io import BytesIO
import tempfile

# ==================== CONFIGURATION ====================
st.set_page_config(
    page_title="EHS Batna - Monitoring Réseau",
    page_icon="🔌",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== CHARGEMENT DES DONNÉES RÉSEAU ====================
def load_data():
    if os.path.exists("reseau_data.csv"):
        df = pd.read_csv("reseau_data.csv")
        colonnes_requises = ["ID", "Side", "Bureau_Num", "Bureau_Nom", "Etage", "Goulotte", "Cable", "Prise", "Statut", "Commentaire", "DerniereModification", "ModifiePar"]
        for col in colonnes_requises:
            if col not in df.columns:
                df[col] = ""
        return df
    
    # Données par défaut (fichier n'existe pas)
    return pd.DataFrame({
        "ID": [1, 2, 3, 4, 5],
        "Side": ["🏛️ Administration", "🏛️ Administration", "🏛️ Administration", "🏥 Medical", "🏥 Medical"],
        "Bureau_Num": ["101", "102", "103", "201", "202"],
        "Bureau_Nom": ["Direction", "Comptabilité", "RH", "Consultation 1", "Consultation 2"],
        "Etage": ["RDC", "RDC", "RDC", "1er", "1er"],
        "Goulotte": ["✅ Terminé", "✅ Terminé", "🏗️ En cours", "✅ Terminé", "⏳ Non commencé"],
        "Cable": ["✅ Tiré", "✅ Tiré", "⏳ Non tiré", "✅ Tiré", "⏳ Non tiré"],
        "Prise": ["✅ Posée", "⏳ Non posée", "⏳ Non posée", "✅ Posée", "⏳ Non posée"],
        "Statut": ["🟢 Terminé", "🟢 Terminé", "🟡 En cours", "🟢 Terminé", "🔴 Non commencé"],
        "Commentaire": ["", "", "Tirage câble cette semaine", "", ""],
        "DerniereModification": ["", "", "", "", ""],
        "ModifiePar": ["", "", "", "", ""]
    })

def save_data(df):
    df.to_csv("reseau_data.csv", index=False)

# ==================== FONCTIONS POUR LE LOGO ====================
def get_logo_base64():
    logo_path = "logo.png"
    if os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None

def display_logo(size=80):
    logo_base64 = get_logo_base64()
    if logo_base64:
        st.markdown(f"""
        <div style="text-align: center; margin-bottom: 20px;">
            <img src="data:image/png;base64,{logo_base64}" width="{size}">
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="text-align: center; margin-bottom: 20px;">
            <div style="font-size: 3rem;">🏥</div>
        </div>
        """, unsafe_allow_html=True)

# ==================== UTILISATEURS ====================
if "users" not in st.session_state:
    st.session_state.users = {
        "djamel": {
            "password": "merzoug2026",
            "nom": "Mr. MERZOUG Djamel",
            "grade": "Ingénieur en Chef",
            "role": "admin",
            "can_manage_members": True
        },
        "tarek": {
            "password": "boureghda2026",
            "nom": "Mr. BOUREGHDA Tarek",
            "grade": "Technicien Supérieur",
            "role": "admin",
            "can_manage_members": True
        }
    }

def authenticate_user(username, password):
    username = username.lower().strip()
    if username in st.session_state.users:
        if st.session_state.users[username]["password"] == password:
            return st.session_state.users[username]
    return None

def add_user(username, password, nom, grade, role):
    st.session_state.users[username.lower()] = {
        "password": password,
        "nom": nom,
        "grade": grade,
        "role": role,
        "can_manage_members": False
    }

def delete_user(username):
    if username in st.session_state.users and username not in ["djamel", "tarek"]:
        del st.session_state.users[username]

# ==================== INITIALISATION ====================
if "df_reseau" not in st.session_state:
    st.session_state.df_reseau = load_data()
if "current_page" not in st.session_state:
    st.session_state.current_page = "Dashboard"
if "theme" not in st.session_state:
    st.session_state.theme = "dark"
if "user_logged_in" not in st.session_state:
    st.session_state.user_logged_in = False
if "user_username" not in st.session_state:
    st.session_state.user_username = None
if "user_nom" not in st.session_state:
    st.session_state.user_nom = None
if "user_grade" not in st.session_state:
    st.session_state.user_grade = None
if "user_role" not in st.session_state:
    st.session_state.user_role = None
if "user_can_manage_members" not in st.session_state:
    st.session_state.user_can_manage_members = False

# ==================== FONCTIONS DE PERMISSIONS ====================
def can_edit():
    return st.session_state.user_role in ["admin", "technicien", "superviseur"]

def can_manage_members():
    return st.session_state.user_can_manage_members

# ==================== THÈME ====================
def toggle_theme():
    st.session_state.theme = "light" if st.session_state.theme == "dark" else "dark"
    st.rerun()

# ==================== PAGE DE LOGIN ====================
if not st.session_state.user_logged_in:
    st.markdown("""
        <style>
        .login-box {
            max-width: 400px;
            margin: 100px auto;
            padding: 40px;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            border-radius: 20px;
            border: 1px solid #333;
            text-align: center;
        }
        .login-box h1 { color: #667eea; }
        .stButton > button {
            background: linear-gradient(135deg, #667eea, #764ba2);
            color: white;
            border: none;
            padding: 10px;
            border-radius: 10px;
            font-weight: bold;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="login-box">', unsafe_allow_html=True)
    st.image("https://img.icons8.com/color/96/000000/hospital-3.png", width=80)
    st.title("🏥 EHS Batna")
    st.markdown("### Monitoring Réseau")
    
    username = st.text_input("👤 Nom d'utilisateur", placeholder="Entrez votre nom d'utilisateur")
    password = st.text_input("🔒 Mot de passe", type="password", placeholder="Entrez votre mot de passe")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🔓 Se connecter", type="primary", use_container_width=True):
            user = authenticate_user(username, password)
            if user:
                st.session_state.user_logged_in = True
                st.session_state.user_username = username
                st.session_state.user_nom = user["nom"]
                st.session_state.user_grade = user["grade"]
                st.session_state.user_role = user["role"]
                st.session_state.user_can_manage_members = user.get("can_manage_members", False)
                st.rerun()
            else:
                st.error("❌ Nom d'utilisateur ou mot de passe incorrect")
    
    with col2:
        st.caption("💡 Contactez Djamel ou Tarek pour obtenir vos identifiants")
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ==================== CSS ====================
if st.session_state.theme == "dark":
    theme_css = """
    <style>
    .stApp { background-color: #0a0a0a; color: #ffffff; }
    [data-testid="stSidebar"] { background: linear-gradient(180deg, #0f0f0f 0%, #1a1a1a 100%); border-right: 1px solid #2a2a2a; }
    .stat-card { background: linear-gradient(135deg, #1e1e2e 0%, #1a1a2e 100%); border: 1px solid #2a2a2a; }
    .bureau-card { background: linear-gradient(135deg, #1e1e2e 0%, #1a1a2e 100%); border: 1px solid #2a2a2a; }
    .header { background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); border: 1px solid #2a2a2a; }
    .footer { border-top: 1px solid #2a2a2a; color: #666; }
    </style>
    """
else:
    theme_css = """
    <style>
    .stApp { background-color: #f0f2f6; color: #1a1a2e; }
    [data-testid="stSidebar"] { background: linear-gradient(180deg, #ffffff 0%, #f0f0f0 100%); border-right: 1px solid #ddd; }
    .stat-card { background: #ffffff; border: 1px solid #ddd; }
    .bureau-card { background: #ffffff; border: 1px solid #ddd; }
    .header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border: none; color: white; }
    .footer { border-top: 1px solid #ddd; color: #888; }
    </style>
    """

st.markdown(theme_css, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stat-card {
        border-radius: 15px;
        padding: 20px;
        text-align: center;
        transition: transform 0.2s;
    }
    .stat-card:hover { transform: translateY(-2px); }
    .stat-number {
        font-size: 2.5rem;
        font-weight: bold;
        background: linear-gradient(135deg, #667eea, #764ba2);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .bureau-card {
        border-radius: 12px;
        padding: 16px;
        margin-bottom: 12px;
        transition: all 0.2s;
    }
    .bureau-card:hover { transform: translateX(4px); }
    .progress-container {
        background-color: #2a2a2a;
        border-radius: 10px;
        height: 8px;
        overflow: hidden;
    }
    .progress-fill {
        background: linear-gradient(90deg, #667eea, #764ba2);
        height: 100%;
        border-radius: 10px;
        transition: width 0.5s;
    }
    .header {
        border-radius: 15px;
        padding: 20px;
        margin-bottom: 20px;
    }
    .footer {
        text-align: center;
        padding: 20px;
        margin-top: 30px;
        font-size: 0.8rem;
    }
    </style>
""", unsafe_allow_html=True)

# ==================== SIDEBAR ====================
with st.sidebar:
    display_logo(size=80)
    st.markdown("### 🔌 EHS Batna")
    st.markdown("#### Monitoring Réseau")
    st.markdown("---")
    
    st.markdown(f"**👤 Connecté :**")
    st.markdown(f"**{st.session_state.user_nom}**")
    st.markdown(f"*{st.session_state.user_grade}*")
    if st.session_state.user_role == "admin":
        st.markdown("🔑 *Administrateur*")
    
    st.markdown("---")
    
    col_theme1, col_theme2 = st.columns([1, 3])
    with col_theme1:
        theme_icon = "🌙" if st.session_state.theme == "dark" else "☀️"
        st.markdown(f"### {theme_icon}")
    with col_theme2:
        if st.button("Changer de thème", key="toggle_theme", use_container_width=True):
            toggle_theme()
    
    st.markdown("---")
    
    st.markdown("### 📋 Navigation")
    
    menu_items = {
        "Dashboard": "📊",
        "Bureaux": "🏢",
        "Réseau": "🌐",
        "Rapports": "📄"
    }
    
    if st.session_state.user_can_manage_members:
        menu_items["Équipe"] = "👥"
    
    for item, icon in menu_items.items():
        if st.button(f"{icon} {item}", key=f"nav_{item}", use_container_width=True):
            st.session_state.current_page = item
            st.rerun()
    
    st.markdown("---")
    
    if st.button("🚪 Se déconnecter", use_container_width=True):
        for key in ["user_logged_in", "user_username", "user_nom", "user_grade", "user_role", "user_can_manage_members"]:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()
    
    st.markdown("---")
    st.caption(f"📅 {datetime.now().strftime('%d/%m/%Y %H:%M')}")

# ==================== DONNÉES ====================
df = st.session_state.df_reseau
total_bureaux = len(df)
termines = len(df[df["Statut"] == "🟢 Terminé"])
en_cours = len(df[df["Statut"] == "🟡 En cours"])
non_commences = len(df[df["Statut"] == "🔴 Non commencé"])
progression = (termines / total_bureaux * 100) if total_bureaux > 0 else 0

# ==================== PAGE DASHBOARD ====================
if st.session_state.current_page == "Dashboard":
    st.markdown('<div class="header">', unsafe_allow_html=True)
    st.title("🏥 EHS Batna - Monitoring Réseau")
    st.markdown("### Suivi d'installation des infrastructures")
    st.markdown(f"*Connecté en tant que : {st.session_state.user_nom}*")
    st.markdown('</div>', unsafe_allow_html=True)
    
    side_filter = st.selectbox("🏢 Filtrer par Side", ["Tous", "🏛️ Administration", "🏥 Medical"])
    
    filtered_df = df.copy()
    if side_filter != "Tous":
        filtered_df = filtered_df[filtered_df["Side"] == side_filter]
    
    total_filtered = len(filtered_df)
    termines_filtered = len(filtered_df[filtered_df["Statut"] == "🟢 Terminé"])
    en_cours_filtered = len(filtered_df[filtered_df["Statut"] == "🟡 En cours"])
    non_commences_filtered = len(filtered_df[filtered_df["Statut"] == "🔴 Non commencé"])
    progression_filtered = (termines_filtered / total_filtered * 100) if total_filtered > 0 else 0
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
        <div class="stat-card">
            <div style="font-size: 2rem;">📊</div>
            <div class="stat-number">{total_filtered}</div>
            <div style="opacity: 0.8;">Bureaux</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="stat-card">
            <div style="font-size: 2rem;">✅</div>
            <div class="stat-number" style="color: #10b981; -webkit-text-fill-color: #10b981;">{termines_filtered}</div>
            <div style="opacity: 0.8;">Terminés</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown(f"""
        <div class="stat-card">
            <div style="font-size: 2rem;">🟡</div>
            <div class="stat-number" style="color: #f59e0b; -webkit-text-fill-color: #f59e0b;">{en_cours_filtered}</div>
            <div style="opacity: 0.8;">En cours</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown(f"""
        <div class="stat-card">
            <div style="font-size: 2rem;">🔴</div>
            <div class="stat-number" style="color: #ef4444; -webkit-text-fill-color: #ef4444;">{non_commences_filtered}</div>
            <div style="opacity: 0.8;">À faire</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown(f"""
    <div>
        <div style="display: flex; justify-content: space-between; margin-bottom: 8px;">
            <span>📈 Progression globale ({side_filter if side_filter != 'Tous' else 'Tous les sides'})</span>
            <span><strong>{progression_filtered:.0f}%</strong> ({termines_filtered}/{total_filtered} bureaux)</span>
        </div>
        <div class="progress-container">
            <div class="progress-fill" style="width: {progression_filtered}%;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    st.markdown("### 📋 Derniers bureaux actifs")
    
    for _, row in filtered_df.head(5).iterrows():
        if row['Statut'] == "🟢 Terminé":
            badge = "✅ Terminé"
        elif row['Statut'] == "🟡 En cours":
            badge = "🟡 En cours"
        else:
            badge = "🔴 Non commencé"
        
        st.markdown(f"""
        <div class="bureau-card">
            <div style="display: flex; justify-content: space-between; align-items: center;">
                <div>
                    <strong>{row['Side']} - {row['Bureau_Num']} - {row['Bureau_Nom']}</strong>
                    <br>
                    <span style="font-size: 0.85rem; opacity: 0.7;">Étage {row['Etage']}</span>
                    <br>
                    <span style="font-size: 0.8rem;">Goulotte {row['Goulotte']} | Câble {row['Cable']} | Prise {row['Prise']}</span>
                </div>
                <div><strong>{badge}</strong></div>
            </div>
        </div>
        """, unsafe_allow_html=True)

# ==================== PAGE BUREAUX ====================
elif st.session_state.current_page == "Bureaux":
    st.markdown('<div class="header">', unsafe_allow_html=True)
    st.title("🏢 Gestion des bureaux")
    st.markdown("### Ajouter, modifier ou supprimer des bureaux")
    st.markdown('</div>', unsafe_allow_html=True)
    
    side_filter = st.selectbox("🏢 Filtrer par Side", ["Tous", "🏛️ Administration", "🏥 Medical"])
    
    filtered_df = df.copy()
    if side_filter != "Tous":
        filtered_df = filtered_df[filtered_df["Side"] == side_filter]
    
    filtered_df = filtered_df.reset_index(drop=True)
    
    for idx, row in filtered_df.iterrows():
        bureau_id = row['ID']
        
        if st.session_state.get(f"edit_mode_{bureau_id}") == True:
            with st.container():
                st.markdown(f"""
                <div class="bureau-card" style="border-color: #f59e0b;">
                    <strong>✏️ Modification du bureau {row['Bureau_Num']} - {row['Bureau_Nom']}</strong>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    new_side = st.selectbox("Side", ["🏛️ Administration", "🏥 Medical"], 
                                           index=0 if row['Side'] == "🏛️ Administration" else 1,
                                           key=f"side_edit_{bureau_id}")
                with col2:
                    new_num = st.text_input("Numéro", value=row['Bureau_Num'], key=f"num_edit_{bureau_id}")
                with col3:
                    new_nom = st.text_input("Nom", value=row['Bureau_Nom'], key=f"nom_edit_{bureau_id}")
                with col4:
                    new_etage = st.selectbox("Étage", ["RDC", "1er", "2ème", "3ème"], 
                                            index=["RDC", "1er", "2ème", "3ème"].index(row['Etage']) if row['Etage'] in ["RDC", "1er", "2ème", "3ème"] else 0,
                                            key=f"etage_edit_{bureau_id}")
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("💾 Sauvegarder", key=f"save_{bureau_id}", type="primary"):
                        idx_in_df = df[df["ID"] == bureau_id].index
                        if len(idx_in_df) > 0:
                            original_idx = idx_in_df[0]
                            df.at[original_idx, 'Side'] = str(new_side)
                            df.at[original_idx, 'Bureau_Num'] = str(new_num)
                            df.at[original_idx, 'Bureau_Nom'] = str(new_nom)
                            df.at[original_idx, 'Etage'] = str(new_etage)
                            df.at[original_idx, 'DerniereModification'] = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
                            df.at[original_idx, 'ModifiePar'] = st.session_state.user_username.upper()
                            save_data(df)
                            st.session_state.df_reseau = df
                            st.session_state[f"edit_mode_{bureau_id}"] = False
                            st.success("✅ Bureau modifié")
                            st.rerun()
                        else:
                            st.error("❌ Bureau non trouvé")
                with col_btn2:
                    if st.button("❌ Annuler", key=f"cancel_{bureau_id}"):
                        st.session_state[f"edit_mode_{bureau_id}"] = False
                        st.rerun()
        else:
            st.markdown(f"""
            <div class="bureau-card">
                <div style="display: flex; justify-content: space-between; align-items: center;">
                    <div>
                        <strong>{row['Side']} - {row['Bureau_Num']} - {row['Bureau_Nom']}</strong>
                        <br>
                        <span style="font-size: 0.85rem; opacity: 0.7;">Étage {row['Etage']}</span>
                        <br>
                        <span style="font-size: 0.7rem; opacity: 0.5;">Dernière modif : {row['DerniereModification']} par {row['ModifiePar']}</span>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("✏️ Modifier", key=f"edit_{bureau_id}"):
                    st.session_state[f"edit_mode_{bureau_id}"] = True
                    st.rerun()
            with col_btn2:
                if st.button("🗑️ Supprimer", key=f"delete_{bureau_id}", type="secondary"):
                    idx_in_df = df[df["ID"] == bureau_id].index
                    if len(idx_in_df) > 0:
                        df = df.drop(idx_in_df[0]).reset_index(drop=True)
                        save_data(df)
                        st.session_state.df_reseau = df
                        st.success("🗑️ Bureau supprimé")
                        st.rerun()
        
        st.markdown("---")
    
    # Ajouter un bureau
    st.markdown("### ➕ Ajouter un nouveau bureau")
    
    with st.form(key="add_bureau_form", clear_on_submit=True):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            new_side = st.selectbox("Side", ["🏛️ Administration", "🏥 Medical"], key="add_side")
        with col2:
            new_num = st.text_input("Numéro", placeholder="Ex: 106", key="add_num")
        with col3:
            new_nom = st.text_input("Nom", placeholder="Ex: Bureau Info", key="add_nom")
        with col4:
            new_etage = st.selectbox("Étage", ["RDC", "1er", "2ème", "3ème"], key="add_etage")
        
        if st.form_submit_button("➕ Ajouter", type="primary", use_container_width=True):
            if new_num and new_nom:
                next_id = df["ID"].max() + 1 if len(df) > 0 else 1
                new_row = pd.DataFrame([{
                    "ID": next_id,
                    "Side": str(new_side),
                    "Bureau_Num": str(new_num),
                    "Bureau_Nom": str(new_nom),
                    "Etage": str(new_etage),
                    "Goulotte": "⏳ Non commencé",
                    "Cable": "⏳ Non tiré",
                    "Prise": "⏳ Non posée",
                    "Statut": "🔴 Non commencé",
                    "Commentaire": "",
                    "DerniereModification": datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                    "ModifiePar": st.session_state.user_username.upper()
                }])
                df = pd.concat([df, new_row], ignore_index=True)
                save_data(df)
                st.session_state.df_reseau = df
                st.success(f"✅ Bureau {new_num} - {new_nom} ajouté")
                st.rerun()
            else:
                st.error("❌ Veuillez remplir tous les champs")

# ==================== PAGE RÉSEAU ====================
elif st.session_state.current_page == "Réseau":
    st.markdown('<div class="header">', unsafe_allow_html=True)
    st.title("🌐 Suivi réseau par bureau")
    st.markdown("### Goulotte, Câble, Prise murale")
    st.markdown('</div>', unsafe_allow_html=True)
    
    col_f0, col_f1, col_f2 = st.columns(3)
    with col_f0:
        side_filter = st.selectbox("🏢 Side", ["Tous", "🏛️ Administration", "🏥 Medical"])
    with col_f1:
        etage_filter = st.selectbox("🏢 Étage", ["Tous"] + sorted(df["Etage"].unique().tolist()))
    with col_f2:
        statut_filter = st.selectbox("📌 Statut", ["Tous", "🟢 Terminé", "🟡 En cours", "🔴 Non commencé"])
    
    filtered_df = df.copy()
    if side_filter != "Tous":
        filtered_df = filtered_df[filtered_df["Side"] == side_filter]
    if etage_filter != "Tous":
        filtered_df = filtered_df[filtered_df["Etage"] == etage_filter]
    if statut_filter != "Tous":
        filtered_df = filtered_df[filtered_df["Statut"] == statut_filter]
    
    for idx, row in filtered_df.iterrows():
        bureau_key = f"{row['Side']}_{row['Bureau_Num']}_{row['Bureau_Nom']}"
        
        if st.session_state.get(f"edit_reseau_{bureau_key}") == True:
            with st.container():
                st.markdown(f"""
                <div class="bureau-card" style="border-color: #f59e0b;">
                    <strong>✏️ Modification réseau - {row['Side']} - {row['Bureau_Num']} {row['Bureau_Nom']}</strong>
                </div>
                """, unsafe_allow_html=True)
                
                col_e1, col_e2 = st.columns(2)
                with col_e1:
                    new_goulotte = st.selectbox("Goulotte", ["⏳ Non commencé", "🏗️ En cours", "✅ Terminé"], 
                                               index=["⏳ Non commencé", "🏗️ En cours", "✅ Terminé"].index(row['Goulotte']),
                                               key=f"goulotte_{bureau_key}")
                    new_cable = st.selectbox("Câble", ["⏳ Non tiré", "🏗️ En cours", "✅ Tiré"],
                                            index=["⏳ Non tiré", "🏗️ En cours", "✅ Tiré"].index(row['Cable']),
                                            key=f"cable_{bureau_key}")
                with col_e2:
                    new_prise = st.selectbox("Prise", ["⏳ Non posée", "🏗️ En cours", "✅ Posée"],
                                            index=["⏳ Non posée", "🏗️ En cours", "✅ Posée"].index(row['Prise']),
                                            key=f"prise_{bureau_key}")
                    new_commentaire = st.text_area("Commentaire", value=row['Commentaire'], key=f"comm_{bureau_key}")
                
                if new_goulotte == "✅ Terminé" and new_cable == "✅ Tiré" and new_prise == "✅ Posée":
                    new_statut = "🟢 Terminé"
                elif new_goulotte != "⏳ Non commencé" or new_cable != "⏳ Non tiré" or new_prise != "⏳ Non posée":
                    new_statut = "🟡 En cours"
                else:
                    new_statut = "🔴 Non commencé"
                
                st.info(f"**Statut calculé :** {new_statut}")
                
                col_b1, col_b2 = st.columns(2)
                with col_b1:
                    if st.button("💾 Sauvegarder", key=f"save_reseau_{bureau_key}", type="primary"):
                        original_mask = (df["Side"] == row["Side"]) & (df["Bureau_Num"] == row["Bureau_Num"]) & (df["Bureau_Nom"] == row["Bureau_Nom"])
                        original_indices = df[original_mask].index.tolist()
                        if original_indices:
                            current_idx = original_indices[0]
                            df.at[current_idx, 'Goulotte'] = new_goulotte
                            df.at[current_idx, 'Cable'] = new_cable
                            df.at[current_idx, 'Prise'] = new_prise
                            df.at[current_idx, 'Statut'] = new_statut
                            df.at[current_idx, 'Commentaire'] = new_commentaire
                            df.at[current_idx, 'DerniereModification'] = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
                            df.at[current_idx, 'ModifiePar'] = st.session_state.user_username.upper()
                            save_data(df)
                            st.session_state.df_reseau = df
                            st.session_state[f"edit_reseau_{bureau_key}"] = False
                            st.success("✅ Réseau modifié")
                            st.rerun()
                with col_b2:
                    if st.button("❌ Annuler", key=f"cancel_reseau_{bureau_key}"):
                        st.session_state[f"edit_reseau_{bureau_key}"] = False
                        st.rerun()
        else:
            if row['Statut'] == "🟢 Terminé":
                badge = "✅ Terminé"
            elif row['Statut'] == "🟡 En cours":
                badge = "🟡 En cours"
            else:
                badge = "🔴 Non commencé"
            
            st.markdown(f"""
            <div class="bureau-card">
                <div style="display: flex; justify-content: space-between; align-items: flex-start;">
                    <div style="flex: 3;">
                        <strong>{row['Side']} - {row['Bureau_Num']} - {row['Bureau_Nom']}</strong>
                        <span style="margin-left: 12px; font-size: 0.85rem; opacity: 0.7;">Étage {row['Etage']}</span>
                        <div style="margin-top: 8px;">
                            <span style="display: inline-block; width: 80px;">Goulotte</span> {row['Goulotte']}<br>
                            <span style="display: inline-block; width: 80px;">Câble</span> {row['Cable']}<br>
                            <span style="display: inline-block; width: 80px;">Prise</span> {row['Prise']}
                        </div>
                        <div style="margin-top: 8px; font-size: 0.85rem; opacity: 0.7;">
                            📝 {row['Commentaire'] if row['Commentaire'] else "Aucun commentaire"}
                        </div>
                        <div style="margin-top: 8px; font-size: 0.7rem; opacity: 0.5;">
                            🕐 Modifié le {row['DerniereModification']} par {row['ModifiePar']}
                        </div>
                    </div>
                    <div style="flex: 1; text-align: right;">
                        <div style="margin-bottom: 8px;"><strong>{badge}</strong></div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("✏️ Modifier l'avancement", key=f"edit_reseau_btn_{bureau_key}"):
                st.session_state[f"edit_reseau_{bureau_key}"] = True
                st.rerun()
        
        st.markdown("---")

# ==================== PAGE ÉQUIPE ====================
elif st.session_state.current_page == "Équipe" and st.session_state.user_can_manage_members:
    st.markdown('<div class="header">', unsafe_allow_html=True)
    st.title("👥 Gestion de l'équipe")
    st.markdown("### Ajouter ou supprimer des membres")
    st.markdown("🔐 *Seuls les administrateurs principaux (Djamel et Tarek) ont accès à cette page*")
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("#### 👤 Membres actuels")
    
    for username, user_data in st.session_state.users.items():
        col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
        with col1:
            st.markdown(f"**{user_data['nom']}**")
        with col2:
            st.markdown(f"*{user_data['grade']}*")
        with col3:
            st.markdown(f"`{username}`")
        with col4:
            if username not in ["djamel", "tarek"]:
                if st.button("🗑️ Supprimer", key=f"del_user_{username}"):
                    delete_user(username)
                    st.toast(f"✅ {user_data['nom']} a été supprimé", icon="🗑️")
                    st.rerun()
        st.divider()
    
    st.markdown("---")
    
    st.markdown("#### ➕ Ajouter un membre")
    
    with st.form("add_user_form", clear_on_submit=True):
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            new_username = st.text_input("Nom d'utilisateur", placeholder="Ex: ahmed")
            new_nom = st.text_input("Nom complet", placeholder="Ex: Mr. BENALI Ahmed")
        with col_a2:
            new_password = st.text_input("Mot de passe", type="password", placeholder="Mot de passe")
            new_grade = st.text_input("Grade", placeholder="Ex: Technicien Réseau")
        
        new_role = st.selectbox("Rôle", ["technicien", "superviseur", "observateur"])
        st.caption("📌 Rôles : technicien (peut modifier), superviseur (peut modifier), observateur (lecture seule)")
        
        if st.form_submit_button("➕ Ajouter le membre", type="primary", use_container_width=True):
            if new_username and new_nom and new_password:
                if new_username.lower() in st.session_state.users:
                    st.error("❌ Ce nom d'utilisateur existe déjà")
                else:
                    add_user(new_username, new_password, new_nom, new_grade, new_role)
                    st.toast(f"✅ {new_nom} a rejoint l'équipe !", icon="👥")
                    st.rerun()
            else:
                st.error("❌ Veuillez remplir tous les champs")
    
    st.markdown("---")
    st.info("💡 Les nouveaux membres pourront se connecter avec leur nom d'utilisateur et mot de passe")

# ==================== PAGE RAPPORTS ====================
elif st.session_state.current_page == "Rapports":
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from docx import Document
    from docx.shared import Inches
    
    st.markdown('<div class="header">', unsafe_allow_html=True)
    st.title("📄 Rapports")
    st.markdown("### Export des données réseau")
    st.markdown('</div>', unsafe_allow_html=True)
    
    logo_base64 = get_logo_base64()
    logo_html = ""
    if logo_base64:
        logo_html = f'<img src="data:image/png;base64,{logo_base64}" width="80" style="margin-bottom: 10px;">'
    
    # ==================== CRÉER UN VRAI PDF ====================
    def create_pdf():
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        tmp_path = None
        if logo_base64:
            try:
                img_data = base64.b64decode(logo_base64)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(img_data)
                    tmp_path = tmp.name
                img = Image(tmp_path, width=50, height=50)
                elements.append(img)
            except Exception:
                pass
        
        elements.append(Spacer(1, 0.2*inch))
        
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, textColor=colors.HexColor('#667eea'))
        elements.append(Paragraph("RAPPORT DE MONITORING RÉSEAU - EHS BATNA", title_style))
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        elements.append(Paragraph("ÉQUIPE TECHNIQUE :", styles['Heading4']))
        elements.append(Paragraph("- Suivi : Mr. MERZOUG Djamel (Ingénieur en Chef)", styles['Normal']))
        elements.append(Paragraph("- Réalisation : Mr. BOUREGHDA Tarek (Technicien Supérieur)", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        elements.append(Paragraph("STATISTIQUES :", styles['Heading4']))
        elements.append(Paragraph(f"- Total bureaux : {total_bureaux}", styles['Normal']))
        elements.append(Paragraph(f"- Terminés : {termines}", styles['Normal']))
        elements.append(Paragraph(f"- En cours : {en_cours}", styles['Normal']))
        elements.append(Paragraph(f"- Non commencés : {non_commences}", styles['Normal']))
        elements.append(Paragraph(f"- Progression : {progression:.0f}%", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        table_data = [["Side", "Bureau", "Nom", "Étage", "Goulotte", "Câble", "Prise", "Statut"]]
        for _, row in df.iterrows():
            table_data.append([
                str(row['Side']), str(row['Bureau_Num']), str(row['Bureau_Nom']), str(row['Etage']),
                str(row['Goulotte']), str(row['Cable']), str(row['Prise']), str(row['Statut'])
            ])
        
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        return buffer
    
    # ==================== CRÉER UN VRAI WORD ====================
    def create_word():
        doc = Document()
        
        if logo_base64:
            try:
                img_data = base64.b64decode(logo_base64)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(img_data)
                    tmp_path = tmp.name
                doc.add_picture(tmp_path, width=Inches(0.8))
                os.unlink(tmp_path)
            except Exception:
                pass
        
        doc.add_heading('RAPPORT DE MONITORING RÉSEAU - EHS BATNA', 0)
        doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph("")
        
        doc.add_heading('ÉQUIPE TECHNIQUE :', level=1)
        doc.add_paragraph('Suivi : Mr. MERZOUG Djamel (Ingénieur en Chef)')
        doc.add_paragraph('Réalisation : Mr. BOUREGHDA Tarek (Technicien Supérieur)')
        doc.add_paragraph("")
        
        doc.add_heading('STATISTIQUES :', level=1)
        doc.add_paragraph(f'Total bureaux : {total_bureaux}')
        doc.add_paragraph(f'Terminés : {termines}')
        doc.add_paragraph(f'En cours : {en_cours}')
        doc.add_paragraph(f'Non commencés : {non_commences}')
        doc.add_paragraph(f'Progression : {progression:.0f}%')
        doc.add_paragraph("")
        
        doc.add_heading('DÉTAIL PAR BUREAU :', level=1)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Side', 'Bureau', 'Nom', 'Étage', 'Goulotte', 'Câble', 'Prise', 'Statut']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Side'])
            row_cells[1].text = str(row['Bureau_Num'])
            row_cells[2].text = str(row['Bureau_Nom'])
            row_cells[3].text = str(row['Etage'])
            row_cells[4].text = str(row['Goulotte'])
            row_cells[5].text = str(row['Cable'])
            row_cells[6].text = str(row['Prise'])
            row_cells[7].text = str(row['Statut'])
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    # ==================== RAPPORT TEXTE ====================
    rapport_txt = f"""
{logo_html}
============================================================
           RAPPORT DE MONITORING RÉSEAU - EHS BATNA
============================================================
Date : {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
Généré par : {st.session_state.user_nom}

ÉQUIPE TECHNIQUE
----------------
Suivi : Mr. MERZOUG Djamel (Ingénieur en Chef)
Réalisation : Mr. BOUREGHDA Tarek (Technicien Supérieur)

STATISTIQUES
------------
Total bureaux : {total_bureaux}
Terminés : {termines}
En cours : {en_cours}
Non commencés : {non_commences}
Progression : {progression:.0f}%

DÉTAIL PAR BUREAU
-----------------
{df.to_string(index=False)}

============================================================
FIN DU RAPPORT
"""
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("📥 Télécharger TXT", rapport_txt, f"rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", use_container_width=True)
    with col2:
        try:
            pdf_buffer = create_pdf()
            st.download_button("📥 Télécharger PDF", pdf_buffer, f"rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", mime="application/pdf", use_container_width=True)
        except Exception as e:
            st.error(f"Erreur PDF : {e}")
    with col3:
        try:
            word_buffer = create_word()
            st.download_button("📥 Télécharger Word", word_buffer, f"rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", use_container_width=True)
        except Exception as e:
            st.error(f"Erreur Word : {e}")
    
    st.markdown("---")
    st.info("💡 Les fichiers PDF et Word sont de vrais documents ouvrables normalement.")

# ==================== FOOTER ====================
st.markdown(f"""
<div class="footer">
    EHS Batna - Service Informatique | Monitoring Réseau v6.0 | Mode {st.session_state.theme.upper()}
    <br>
    Connecté : {st.session_state.user_nom}
</div>
""", unsafe_allow_html=True)
